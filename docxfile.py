import docx
import os
os.chdir('F:\python pgrms')
doc=docx.Document()
doc.add_heading("heading for the document",2)
doc_para=doc.add_paragraph("this is paragraph")
doc_para.add_run("this is sindhu").bold
doc.save("docx.docx")
